from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "forms"
OUT_PATH = OUT_DIR / "李方舟法定代理人同意書.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, size=11, bold=False, color="000000"):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def add_labeled_table(doc, rows, widths=(2520, 6840)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, list(widths))
    for row_idx, (label, value) in enumerate(rows):
        label_cell = table.cell(row_idx, 0)
        value_cell = table.cell(row_idx, 1)
        set_cell_shading(label_cell, "F2F4F7")
        label_p = label_cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        add_text(label_p, label, bold=True, color="1F4D78")
        value_p = value_cell.paragraphs[0]
        value_p.paragraph_format.space_after = Pt(0)
        add_text(value_p, value)
    return table


def blank(length=34):
    return "＿" * length


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    add_text(title, "2026 和泰 AI 黑客松", size=15, bold=True, color="2E74B5")
    title.add_run("\n")
    add_text(title, "法定代理人同意書", size=20, bold=True, color="0B2545")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(12)
    add_text(meta, "隊伍：iRent Guard｜AI 車況守護系統　　隊長：侯冠宇")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    add_text(
        note,
        "說明：本文件依競賽報名頁所載「未滿 18 歲參賽者須檢附法定代理人同意書」需求製作。若主辦單位提供指定格式，請以主辦單位指定格式為準。",
        size=10,
        color="555555",
    )

    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(6)
    add_text(h1, "一、參賽者資料", size=13, bold=True, color="2E74B5")
    add_labeled_table(
        doc,
        [
            ("參賽者姓名", "李方舟"),
            ("出生年月日", blank(22)),
            ("身分證字號/居留證號", blank(22)),
            ("就讀學校/單位", blank(26)),
            ("電子信箱", "d1457368464@gmail.com"),
            ("聯絡電話", blank(26)),
        ],
    )

    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    add_text(h2, "二、同意事項", size=13, bold=True, color="2E74B5")

    consent = [
        "本人為上列未成年參賽者之法定代理人，已知悉並同意其以隊伍成員身分參加「2026 和泰 AI 黑客松」。",
        "本人同意參賽者配合競賽報名、初賽提案繳交、線上或實體活動、工作坊、簡報與成果展示等必要流程。",
        "本人已提醒參賽者遵守競賽規則、主辦單位通知、智慧財產權與個人資料相關規範，並願配合主辦單位辦理必要之資格審查。",
        "若隊伍晉級決賽或主辦單位另有指定授權、肖像或作品相關文件，本人同意依主辦單位要求另行審閱並簽署。",
    ]
    for idx, item in enumerate(consent, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        add_text(p, f"{idx}. ", bold=True)
        add_text(p, item)

    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    add_text(h3, "三、法定代理人資料", size=13, bold=True, color="2E74B5")
    add_labeled_table(
        doc,
        [
            ("法定代理人姓名", blank(26)),
            ("與參賽者關係", "□ 父　□ 母　□ 監護人　□ 其他：" + blank(12)),
            ("身分證字號/居留證號", blank(22)),
            ("聯絡電話", blank(26)),
            ("電子信箱", blank(26)),
            ("通訊地址", blank(30)),
        ],
    )

    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    add_text(h4, "四、簽署", size=13, bold=True, color="2E74B5")

    sig = doc.add_table(rows=4, cols=2)
    sig.style = "Table Grid"
    set_table_geometry(sig, [4680, 4680])
    signature_rows = [
        ("法定代理人簽章", blank(18)),
        ("參賽者簽名", blank(18)),
        ("簽署日期", "中華民國　　　　年　　　　月　　　　日"),
        ("備註", "本同意書完成簽署後，請掃描或拍照上傳至官方報名系統。"),
    ]
    for row_idx, (label, value) in enumerate(signature_rows):
        for cell in sig.rows[row_idx].cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(sig.cell(row_idx, 0), "F2F4F7")
        add_text(sig.cell(row_idx, 0).paragraphs[0], label, bold=True, color="1F4D78")
        add_text(sig.cell(row_idx, 1).paragraphs[0], value)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "iRent Guard｜AI 車況守護系統｜2026 和泰 AI 黑客松", size=9, color="555555")

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
